import streamlit as st
import os
import subprocess
from pydub import AudioSegment
import speech_recognition as sr
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from pdf2image import convert_from_path
import tempfile

# Set the FFMPEG_BINARY to the direct path of ffmpeg executable
os.environ["FFMPEG_BINARY"] = "/usr/local/bin/ffmpeg"
os.environ["FFPROBE_BINARY"] = "/usr/local/bin/ffprobe"

# Function to check if ffmpeg is accessible
def check_ffmpeg():
    try:
        # Run ffmpeg command to check its version
        result = subprocess.run(['sudo', 'ffmpeg', '-version'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        st.write("FFmpeg version detected successfully:")
        st.write(result.stdout)
    except FileNotFoundError:
        st.error("FFmpeg is not installed or not found in the specified path.")
    except PermissionError:
        st.error("Permission denied while trying to run ffmpeg with sudo.")
    except Exception as e:
        st.error(f"An error occurred: {e}")

# Function 1: Play "engineer_diagnosis.wav" file from local directory
def play_engineer_diagnosis():
    st.header("Stage 1: Play 'engineer_diagnosis.wav'")
    file_path = "engineer_diagnosis.wav"
    if os.path.exists(file_path):
        st.audio(file_path, format="audio/wav")
    else:
        st.error("File 'engineer_diagnosis.wav' not found!")

# Function 2: Record voice, save as wav, and allow playback
def record_voice():
    st.header("Stage 2: Record Voice")
    uploaded_file = st.file_uploader("Choose a file", type=["wav", "mp3"])
    if uploaded_file:
        audio = AudioSegment.from_file(uploaded_file)
        wav_file_path = "uploaded_audio.wav"
        audio.export(wav_file_path, format="wav")
        st.success(f"Audio uploaded and saved as {wav_file_path}!")
        st.audio(wav_file_path, format="audio/wav")

# Function 3: Convert speech from "electric_unit_heater.wav" file to text using Google Speech Recognition
def process_speech_to_text():
    st.header("Stage 3: Recognize Speech from 'electric_unit_heater.wav'")
    wav_file = "electric_unit_heater.wav"
    if os.path.exists(wav_file):
        recognizer = sr.Recognizer()
        with sr.AudioFile(wav_file) as source:
            audio = recognizer.record(source)
            try:
                text = recognizer.recognize_google(audio)
                st.write("Recognized Text:")
                sentences = text.split('.')
                for i, sentence in enumerate(sentences, 1):
                    st.write(f"{i}. {sentence.strip()}")
                return sentences
            except sr.UnknownValueError:
                st.error("Speech not recognized.")
            except sr.RequestError as e:
                st.error(f"Error with the request: {e}")
    else:
        st.error(f"File '{wav_file}' not found!")
    return []

# Function 4: Create a checklist from recognized speech
def create_checklist_document(sentences):
    st.header("Stage 4: Create Checklist from Recognized Speech")
    if not sentences:
        st.error("No sentences provided to generate the checklist.")
        return

    # Create document
    document = Document()
    document.add_heading('Field Engineer Checklist', level=1)

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Question'
    hdr_cells[1].text = 'Answer (YES/NO)'

    for sentence in sentences:
        if sentence.strip():
            row_cells = table.add_row().cells
            row_cells[0].text = sentence.strip()
            row_cells[1].text = '☐ YES ☐ NO'

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Checklist Document",
        data=buffer,
        file_name="field_engineer_checklist.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Call the function to display the first page as a visual preview
    display_docx_as_image(buffer)

# Convert .docx to PDF and then PDF to image (first page)
def convert_docx_to_pdf(buffer):
    temp_pdf_path = tempfile.mktemp(suffix=".pdf")
    document = Document(buffer)
    c = canvas.Canvas(temp_pdf_path, pagesize=letter)
    c.setFont("Helvetica", 10)
    for paragraph in document.paragraphs:
        c.drawString(100, 750, paragraph.text)
        c.showPage()
    c.save()
    return temp_pdf_path

def display_docx_as_image(buffer):
    try:
        # Convert the .docx to PDF
        temp_pdf_path = convert_docx_to_pdf(buffer)

        # Convert the first page of the PDF to an image
        images = convert_from_path(temp_pdf_path, first_page=1, last_page=1)

        # Display the first page image in Streamlit
        st.image(images[0], caption="First Page of the Checklist Document", use_column_width=True)

    except Exception as e:
        st.error(f"An error occurred while converting or displaying the document: {e}")

# Main App
st.title("Virtual Verbal Assisstant")

if st.button("Stage 1: Play 'engineer_diagnosis.wav'"):
    play_engineer_diagnosis()

if st.button("Stage 2: Record Voice"):
    record_voice()

# Trigger Stage 4 automatically after Stage 3 (if text is recognized)
if st.button("Stage 3: Recognize Speech from 'electric_unit_heater.wav'"):
    sentences = process_speech_to_text()
    if sentences:
        create_checklist_document(sentences)  # Automatically trigger Stage 4 once text is recognized
